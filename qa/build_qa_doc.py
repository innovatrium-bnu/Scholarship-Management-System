# -*- coding: utf-8 -*-
"""Generates the QA report as a real .docx."""

import copy
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

INK = RGBColor(0x16, 0x1B, 0x21)
ACCENT = RGBColor(0x14, 0x45, 0x6B)
MUTED = RGBColor(0x5A, 0x64, 0x70)
FAIL = RGBColor(0xA3, 0x20, 0x17)
PASS = RGBColor(0x17, 0x66, 0x3F)
HDR_FILL = "14456B"
ALT_FILL = "F2F5F8"
WARN_FILL = "FBEFDC"
DEF_FILL = "F8E5E3"
FIX_FILL = "E2F0E8"

BODY = "Calibri"
MONO = "Consolas"

doc = Document()

# ---------------------------------------------------------------- page setup
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
sec.left_margin = sec.right_margin = Cm(1.5)
sec.top_margin = Cm(1.5)
sec.bottom_margin = Cm(1.4)

normal = doc.styles["Normal"]
normal.font.name = BODY
normal.font.size = Pt(10)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for name, size, color, before, after in [
    ("Heading 1", 20, ACCENT, 18, 8),
    ("Heading 2", 14, ACCENT, 16, 6),
    ("Heading 3", 11.5, INK, 12, 4),
]:
    st = doc.styles[name]
    st.font.name = BODY
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True


# ------------------------------------------------------------------- helpers
def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def cell_text(cell, text, *, bold=False, size=9, color=INK, mono=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = MONO if mono else BODY
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return p


def para(text, *, size=10, bold=False, italic=False, color=INK, space_after=6, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(it)
        run.font.size = Pt(10)
        run.font.color.rgb = INK


def info_table(rows, widths, header=None, fills=None):
    """A simple key/value or reference table."""
    cols = len(widths)
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    if header:
        hr = t.add_row()
        repeat_header(hr)
        for i, h in enumerate(header):
            shade(hr.cells[i], HDR_FILL)
            cell_text(hr.cells[i], h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    for n, row in enumerate(rows):
        r = t.add_row()
        for i, val in enumerate(row):
            fill = fills[n] if fills else (ALT_FILL if n % 2 else None)
            if fill:
                shade(r.cells[i], fill)
            cell_text(r.cells[i], str(val), size=9, bold=(i == 0 and not header))
    for r in t.rows:
        for i, w in enumerate(widths):
            r.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ---- the test-case table: ID | Steps | Expected Result | P/F | Notes -------
WIDTHS = [2.0, 9.2, 9.0, 1.5, 5.0]


def test_table(prefix, cases, start=1):
    t = doc.add_table(rows=0, cols=5)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    hr = t.add_row()
    repeat_header(hr)
    for i, h in enumerate(["Test ID", "Steps", "Expected Result", "P/F", "Notes"]):
        shade(hr.cells[i], HDR_FILL)
        cell_text(hr.cells[i], h, bold=True, size=9,
                  color=RGBColor(0xFF, 0xFF, 0xFF),
                  align=WD_ALIGN_PARAGRAPH.CENTER if i == 3 else None)

    n = start
    for case in cases:
        steps, expected = case[0], case[1]
        note = case[2] if len(case) > 2 else ""
        r = t.add_row()
        tid = "%s-%03d" % (prefix, n)
        fill = None
        if note.startswith("FIXED"):
            fill = FIX_FILL
        elif note.startswith("DEFECT"):
            fill = DEF_FILL
        elif note.startswith("KNOWN"):
            fill = WARN_FILL
        elif n % 2 == 0:
            fill = ALT_FILL
        if fill:
            for c in r.cells:
                shade(c, fill)
        cell_text(r.cells[0], tid, size=8.5, mono=True, bold=True, color=ACCENT)
        cell_text(r.cells[1], steps, size=9)
        cell_text(r.cells[2], expected, size=9)
        cell_text(r.cells[3], "", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(r.cells[4], note, size=8.5,
                  color=FAIL if note.startswith("DEFECT")
                  else (PASS if note.startswith("FIXED") else MUTED))
        n += 1

    for r in t.rows:
        for i, w in enumerate(WIDTHS):
            r.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return n


def section(code, title, intro, cases):
    doc.add_heading(title, level=2)
    if intro:
        para(intro, size=9.5, color=MUTED, space_after=8)
    test_table(code, cases)


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ------------------------------------------------------------------ footer
footer_p = sec.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run("BNU Scholarship Management System  ·  QA Test Report v1.0  ·  Page ")
fr.font.size = Pt(8)
fr.font.color.rgb = MUTED
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
footer_p._p.append(fld)

exec(open(__file__.replace("build_qa_doc.py", "qa_content.py"), encoding="utf-8").read())

doc.save(OUT)
print("written:", OUT)
